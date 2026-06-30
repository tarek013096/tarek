import asyncio
import logging
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from telegram import BotCommand, KeyboardButton, MenuButtonCommands, ReplyKeyboardMarkup, Update
from telegram.constants import ChatAction
from telegram.error import TelegramError
from telegram.ext import (
    Application,
    CommandHandler,
    ContextTypes,
    ConversationHandler,
    MessageHandler,
    filters,
)


logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
LOGGER = logging.getLogger(__name__)

BOT_TOKEN = os.environ.get("BOT_TOKEN")
WEBHOOK_URL = os.environ.get("WEBHOOK_URL", "").rstrip("/")
PORT = int(os.environ.get("PORT", "10000"))

LAB_BUTTON = "Lab Report Generator"
ASSIGNMENT_BUTTON = "Assignment Generator"
BACK_BUTTON = "Back"
CLEAR_BUTTON = "Clear"
CANCEL_BUTTON = "Cancel"

MAIN_KEYBOARD = ReplyKeyboardMarkup(
    [
        [KeyboardButton(LAB_BUTTON), KeyboardButton(ASSIGNMENT_BUTTON)],
        [KeyboardButton(CLEAR_BUTTON)],
    ],
    resize_keyboard=True,
    one_time_keyboard=False,
    is_persistent=True,
)

FORM_KEYBOARD = ReplyKeyboardMarkup(
    [
        [KeyboardButton(BACK_BUTTON), KeyboardButton(CLEAR_BUTTON)],
        [KeyboardButton(CANCEL_BUTTON)],
    ],
    resize_keyboard=True,
    one_time_keyboard=False,
    is_persistent=True,
)

FIELDS = [
    ("university_name", "University name likhun:"),
    ("course_code", "Course code likhun:"),
    ("course_title", "Course title likhun:"),
    ("work_no", "Assignment/Report no likhun:"),
    ("student_name", "Student er name likhun:"),
    ("student_id", "Student ID likhun:"),
    ("program", "Program likhun:"),
    ("batch_year_semester", "Batch Year Semester likhun, example: 1st Batch 2nd Year 2nd Semester"),
    ("teacher_name", "Teacher er name likhun:"),
    ("teacher_designation", "Teacher er designation likhun:"),
    ("teacher_department", "Teacher er department likhun:"),
    ("submission_date", "Submission date likhun, example: June 29, 2026"),
]

ASK_FIELD, ASK_LOGO = range(2)
ORDINAL_PATTERN = re.compile(r"(?i)(\d+)(st|nd|rd|th)\b")


def clean_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_. -]+", "", value).strip()
    return value[:70] or "cover-page"


def fit_pdf_text(pdf: canvas.Canvas, text: str, font: str, size: int, max_width: float | None) -> str:
    if not max_width or pdf.stringWidth(text, font, size) <= max_width:
        return text

    suffix = "..."
    fitted = text
    while fitted and pdf.stringWidth(f"{fitted}{suffix}", font, size) > max_width:
        fitted = fitted[:-1].rstrip()
    return f"{fitted}{suffix}" if fitted else suffix


def draw_pdf_text_with_ordinals(
    pdf: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    size: int,
    color,
    font: str = "Helvetica-Bold",
    max_width: float | None = None,
) -> None:
    text = fit_pdf_text(pdf, text, font, size, max_width)
    cursor_x = x
    position = 0
    superscript_size = max(7, int(size * 0.62))
    superscript_y = y + (size * 0.36)

    pdf.setFillColor(color)
    for match in ORDINAL_PATTERN.finditer(text):
        prefix = text[position : match.start()]
        number, suffix = match.groups()

        if prefix:
            pdf.setFont(font, size)
            pdf.drawString(cursor_x, y, prefix)
            cursor_x += pdf.stringWidth(prefix, font, size)

        pdf.setFont(font, size)
        pdf.drawString(cursor_x, y, number)
        cursor_x += pdf.stringWidth(number, font, size)

        pdf.setFont(font, superscript_size)
        pdf.drawString(cursor_x, superscript_y, suffix.lower())
        cursor_x += pdf.stringWidth(suffix.lower(), font, superscript_size)
        position = match.end()

    trailing = text[position:]
    if trailing:
        pdf.setFont(font, size)
        pdf.drawString(cursor_x, y, trailing)


def remember_message(context: ContextTypes.DEFAULT_TYPE, message_id: int | None) -> None:
    if not message_id:
        return
    first_id = context.chat_data.get("first_message_id")
    if first_id is None or message_id < first_id:
        context.chat_data["first_message_id"] = message_id
    ids = context.chat_data.setdefault("message_ids", [])
    if message_id not in ids:
        ids.append(message_id)


async def send_tracked(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str, reply_markup=None):
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)
    await asyncio.sleep(0.35)
    message = await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=text,
        reply_markup=reply_markup,
    )
    remember_message(context, message.message_id)
    return message


async def track_incoming(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if update.effective_message:
        remember_message(context, update.effective_message.message_id)


async def delete_tracked_messages(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    chat_id = update.effective_chat.id
    message_ids = set(context.chat_data.get("message_ids", []))
    first_id = context.chat_data.get("first_message_id")
    current_id = update.effective_message.message_id if update.effective_message else None
    if first_id and current_id and current_id >= first_id:
        message_ids.update(range(first_id, current_id + 1))

    context.chat_data["message_ids"] = []
    context.chat_data.pop("first_message_id", None)

    for message_id in sorted(message_ids, reverse=True):
        try:
            await context.bot.delete_message(chat_id=chat_id, message_id=message_id)
        except TelegramError:
            pass


def set_document_background(document: Document, color: str = "F7FBFC") -> None:
    background = document._element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        document._element.insert(0, background)
    background.set(qn("w:color"), color)


def set_page_border(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)

    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "thinThickThinMediumGap")
        border.set(qn("w:sz"), "22")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), "167A8F")
        borders.append(border)
    sect_pr.append(borders)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "167A8F", size: str = "14", style: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for side in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), style)
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def clear_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for side in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def set_paragraph_bottom_border(paragraph, color: str = "167A8F") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def add_text(
    paragraph,
    text: str,
    size: int,
    bold: bool = False,
    color: RGBColor | None = None,
    font: str = "Aptos Display",
):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_centered_text(
    document: Document,
    text: str,
    size: int,
    bold: bool = False,
    color: RGBColor | None = None,
    space_after: int = 6,
    font: str = "Aptos Display",
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    add_text(paragraph, text, size=size, bold=bold, color=color, font=font)
    return paragraph


def add_spacer(document: Document, points: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def add_course_line(document: Document, course_code: str, course_title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.right_indent = Cm(1.0)
    paragraph.paragraph_format.space_after = Pt(5)
    add_text(paragraph, "COURSE CODE: ", 13, True, RGBColor(0, 128, 64))
    add_text(paragraph, course_code, 13, True, RGBColor(20, 20, 20), "Aptos")

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.left_indent = Cm(1.0)
    title_paragraph.paragraph_format.right_indent = Cm(1.0)
    title_paragraph.paragraph_format.space_after = Pt(20)
    add_text(title_paragraph, "COURSE TITLE: ", 13, True, RGBColor(0, 128, 64))
    add_text(title_paragraph, course_title, 13, True, RGBColor(20, 20, 20), "Aptos")
    set_paragraph_bottom_border(title_paragraph)


def add_label_value(paragraph, label: str, value: str, show_label: bool = True) -> None:
    if show_label:
        add_text(paragraph, label, 10, True, RGBColor(0, 128, 64), "Aptos")
    add_text(paragraph, value, 10, True, RGBColor(22, 31, 38), "Aptos")


def fill_info_box(cell, heading: str, items: list[tuple[str, str, bool]]) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, "F1FAF7")
    set_cell_border(cell, "167A8F", "16", "single")

    heading_paragraph = cell.paragraphs[0]
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_paragraph.paragraph_format.space_after = Pt(7)
    add_text(heading_paragraph, heading, 14, True, RGBColor(0, 128, 64), "Aptos Display")

    for label, value, show_label in items:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(0.12)
        paragraph.paragraph_format.right_indent = Cm(0.12)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
        add_label_value(paragraph, label, value, show_label=show_label)


def add_submission_boxes(document: Document, data: dict) -> None:
    wrapper = document.add_table(rows=1, cols=3)
    wrapper.autofit = False
    left_margin_cell = wrapper.cell(0, 0)
    content_cell = wrapper.cell(0, 1)
    right_margin_cell = wrapper.cell(0, 2)
    for cell, width in ((left_margin_cell, 350), (content_cell, 8350), (right_margin_cell, 350)):
        cell.text = ""
        set_cell_width(cell, width)
        clear_cell_border(cell)
        set_cell_shading(cell, "F7FBFC")
        if cell.paragraphs:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)

    table = content_cell.add_table(rows=1, cols=3)
    table.autofit = False
    left_cell = table.cell(0, 0)
    gap_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    for cell, width in ((left_cell, 3900), (gap_cell, 550), (right_cell, 3900)):
        cell.text = ""
        set_cell_width(cell, width)

    clear_cell_border(gap_cell)
    set_cell_shading(gap_cell, "F7FBFC")

    submitted_by = [
        ("NAME: ", data["student_name"], True),
        ("ID: ", data["student_id"], True),
        ("PROGRAM: ", data["program"], True),
        ("", data["batch_year_semester"], False),
    ]
    submitted_to = [
        ("NAME: ", data["teacher_name"], True),
        ("DESIGNATION: ", data["teacher_designation"], True),
        ("DEPARTMENT: ", data["teacher_department"], True),
    ]

    fill_info_box(left_cell, "SUBMITTED BY:", submitted_by)
    fill_info_box(right_cell, "SUBMITTED TO:", submitted_to)


def create_cover_page(data: dict, logo_path: Path | None) -> Path:
    output_dir = Path(tempfile.gettempdir()) / "cover_page_bot"
    output_dir.mkdir(parents=True, exist_ok=True)
    title = "LAB REPORT" if data["document_type"] == "lab" else "ASSIGNMENT"
    base_name = clean_filename(f"{title.lower()}_{data['student_id']}_{data['work_no']}")
    output_path = output_dir / f"{base_name}.pdf"

    page_width, page_height = A4
    pdf = canvas.Canvas(str(output_path), pagesize=A4)

    green = HexColor("#00A651")
    teal = HexColor("#167A8F")
    dark = HexColor("#15222A")
    light_bg = HexColor("#F7FBFC")
    soft_box = HexColor("#F1FAF7")

    pdf.setFillColor(light_bg)
    pdf.rect(0, 0, page_width, page_height, stroke=0, fill=1)

    for offset, line_width in ((26, 2.2), (32, 0.8)):
        pdf.setStrokeColor(teal)
        pdf.setLineWidth(line_width)
        pdf.rect(offset, offset, page_width - (offset * 2), page_height - (offset * 2), stroke=1, fill=0)

    def center_text(text: str, y: float, size: int, color=dark, font="Helvetica-Bold", max_width: float | None = None) -> None:
        fitted_size = size
        if max_width:
            while fitted_size > 9 and pdf.stringWidth(text, font, fitted_size) > max_width:
                fitted_size -= 1
        pdf.setFont(font, fitted_size)
        pdf.setFillColor(color)
        pdf.drawCentredString(page_width / 2, y, text)

    center_text(data["university_name"].upper(), 778, 28, green, max_width=page_width - 70)

    if logo_path and logo_path.exists():
        try:
            image = ImageReader(str(logo_path))
            logo_size = 112
            pdf.drawImage(
                image,
                (page_width - logo_size) / 2,
                646,
                width=logo_size,
                height=logo_size,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            LOGGER.exception("Could not insert uploaded logo into PDF")

    center_text(title, 603, 32, teal)
    center_text(f"{title.title()} No: {data['work_no']}", 571, 18, dark, max_width=page_width - 120)

    center_text(f"COURSE CODE: {data['course_code']}", 532, 16, dark, max_width=page_width - 120)
    center_text(f"COURSE TITLE: {data['course_title']}", 506, 16, dark, max_width=page_width - 120)
    pdf.setStrokeColor(teal)
    pdf.setLineWidth(1.6)
    pdf.line(95, 488, page_width - 95, 488)

    box_width = 250
    box_height = 190
    gap = 20
    total_width = (box_width * 2) + gap
    left_x = (page_width - total_width) / 2
    right_x = left_x + box_width + gap
    box_y = 260

    def draw_box(x: float, heading: str, rows: list[tuple[str, str, bool]]) -> None:
        pdf.setFillColor(soft_box)
        pdf.setStrokeColor(teal)
        pdf.setLineWidth(1.6)
        pdf.roundRect(x, box_y, box_width, box_height, 9, stroke=1, fill=1)

        pdf.setFillColor(green)
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawCentredString(x + (box_width / 2), box_y + box_height - 34, heading)

        y = box_y + box_height - 70
        for label, value, show_label in rows:
            pdf.setFillColor(green)
            pdf.setFont("Helvetica-Bold", 13)
            if show_label:
                pdf.drawString(x + 20, y, label)
                label_width = pdf.stringWidth(label, "Helvetica-Bold", 13)
                value_x = x + 20 + label_width + 2
            else:
                value_x = x + 20

            pdf.setFillColor(dark)
            draw_pdf_text_with_ordinals(
                pdf,
                value,
                value_x,
                y,
                13,
                dark,
                "Helvetica-Bold",
                max_width=(x + box_width - 20) - value_x,
            )
            y -= 31

    draw_box(
        left_x,
        "SUBMITTED BY:",
        [
            ("NAME: ", data["student_name"], True),
            ("ID: ", data["student_id"], True),
            ("PROGRAM: ", data["program"], True),
            ("", data["batch_year_semester"], False),
        ],
    )
    draw_box(
        right_x,
        "SUBMITTED TO:",
        [
            ("NAME: ", data["teacher_name"], True),
            ("DESIGNATION: ", data["teacher_designation"], True),
            ("DEPARTMENT: ", data["teacher_department"], True),
        ],
    )

    center_text("SUBMISSION DATE:", 150, 17, green)
    center_text(data["submission_date"], 123, 16, dark, max_width=page_width - 140)

    pdf.showPage()
    pdf.save()
    return output_path


def reset_form(context: ContextTypes.DEFAULT_TYPE, document_type: str | None = None) -> None:
    context.user_data.clear()
    if document_type:
        context.user_data["document_type"] = document_type
        context.user_data["field_index"] = 0


async def show_main_menu(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str) -> int:
    await send_tracked(update, context, text, MAIN_KEYBOARD)
    return ConversationHandler.END


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    reset_form(context)
    return await show_main_menu(
        update,
        context,
        "Cover page bot ready. Menu theke Lab Report ba Assignment choose korun.",
    )


async def clear(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    reset_form(context)
    await delete_tracked_messages(update, context)
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)
    await asyncio.sleep(0.35)
    message = await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text="Cleared. Notun kore start korte menu use korun.",
        reply_markup=MAIN_KEYBOARD,
    )
    remember_message(context, message.message_id)
    return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    reset_form(context)
    return await show_main_menu(update, context, "Generation cancel kora hoyeche. Menu theke abar choose korun.")


async def choose_type(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    text = update.effective_message.text.strip()
    if text == LAB_BUTTON:
        reset_form(context, "lab")
    elif text == ASSIGNMENT_BUTTON:
        reset_form(context, "assignment")
    else:
        return await show_main_menu(update, context, "Menu theke valid option choose korun.")

    await send_tracked(update, context, FIELDS[0][1], FORM_KEYBOARD)
    return ASK_FIELD


async def back_field(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    field_index = int(context.user_data.get("field_index", 0))
    if field_index <= 0:
        await send_tracked(update, context, "Eita prothom step. Menu theke notun option choose korte paren.", FORM_KEYBOARD)
        return ASK_FIELD

    field_index -= 1
    field_name, prompt = FIELDS[field_index]
    context.user_data.pop(field_name, None)
    context.user_data["field_index"] = field_index
    await send_tracked(update, context, f"Back kora holo. {prompt}", FORM_KEYBOARD)
    return ASK_FIELD


async def collect_field(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    text = update.effective_message.text.strip()

    if text in {LAB_BUTTON, ASSIGNMENT_BUTTON}:
        return await choose_type(update, context)

    field_index = int(context.user_data.get("field_index", 0))
    field_name, _ = FIELDS[field_index]

    if not text:
        await send_tracked(update, context, "Ei field ta khali rakha jabe na. Abar likhun:", FORM_KEYBOARD)
        return ASK_FIELD

    context.user_data[field_name] = text
    field_index += 1
    context.user_data["field_index"] = field_index

    if field_index < len(FIELDS):
        await send_tracked(update, context, FIELDS[field_index][1], FORM_KEYBOARD)
        return ASK_FIELD

    await send_tracked(
        update,
        context,
        "University logo upload korun. Photo hisebe ba image document hisebe pathan.",
        FORM_KEYBOARD,
    )
    return ASK_LOGO


async def back_from_logo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    context.user_data["field_index"] = len(FIELDS) - 1
    context.user_data.pop("submission_date", None)
    await send_tracked(update, context, f"Back kora holo. {FIELDS[-1][1]}", FORM_KEYBOARD)
    return ASK_FIELD


async def collect_logo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.UPLOAD_DOCUMENT)

    logo_path = None
    file_id = None
    if update.effective_message.photo:
        file_id = update.effective_message.photo[-1].file_id
        logo_path = Path(tempfile.gettempdir()) / f"logo_{update.effective_user.id}.jpg"
    elif update.effective_message.document and (update.effective_message.document.mime_type or "").startswith("image/"):
        file_id = update.effective_message.document.file_id
        extension = Path(update.effective_message.document.file_name or "logo.png").suffix or ".png"
        logo_path = Path(tempfile.gettempdir()) / f"logo_{update.effective_user.id}{extension}"

    if not file_id or not logo_path:
        await send_tracked(update, context, "Logo hisebe ekta image upload korun. Vul hole Back click korun.", FORM_KEYBOARD)
        return ASK_LOGO

    telegram_file = await context.bot.get_file(file_id)
    await telegram_file.download_to_drive(custom_path=str(logo_path))

    data = dict(context.user_data)
    missing = [field for field, _ in FIELDS if not data.get(field)]
    if missing:
        context.user_data["field_index"] = 0
        await send_tracked(update, context, "Kichu input missing. Notun kore first step theke shuru korun.", FORM_KEYBOARD)
        return ASK_FIELD

    output_path = create_cover_page(data, logo_path)
    with output_path.open("rb") as document_file:
        sent_doc = await context.bot.send_document(
            chat_id=update.effective_chat.id,
            document=document_file,
            filename=output_path.name,
            caption="Cover page ready. PDF file ta download kore nin.",
            reply_markup=MAIN_KEYBOARD,
        )
    remember_message(context, sent_doc.message_id)
    reset_form(context)
    await send_tracked(update, context, "Abar generate korte chaile menu theke choose korun.", MAIN_KEYBOARD)
    return ConversationHandler.END


async def ask_logo_again(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await track_incoming(update, context)
    await send_tracked(update, context, "Logo image upload korun, ba Back diye submission date edit korun.", FORM_KEYBOARD)
    return ASK_LOGO


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await track_incoming(update, context)
    await send_tracked(
        update,
        context,
        "Menu theke generator choose korun. Protiti step-e Back/Clear ache. Clear old bot messages delete korar try korbe.",
        MAIN_KEYBOARD,
    )


async def setup_bot_menu(app: Application) -> None:
    await app.bot.set_my_commands(
        [
            BotCommand("start", "Restart the bot"),
            BotCommand("help", "Show help"),
            BotCommand("cancel", "Cancel current generation"),
        ]
    )
    await app.bot.set_chat_menu_button(menu_button=MenuButtonCommands())


def build_application() -> Application:
    if not BOT_TOKEN:
        raise RuntimeError("BOT_TOKEN environment variable missing.")

    app = Application.builder().token(BOT_TOKEN).post_init(setup_bot_menu).build()
    conversation = ConversationHandler(
        entry_points=[
            MessageHandler(filters.Regex(f"^({re.escape(LAB_BUTTON)}|{re.escape(ASSIGNMENT_BUTTON)})$"), choose_type),
        ],
        states={
            ASK_FIELD: [
                MessageHandler(filters.Regex(f"^{re.escape(BACK_BUTTON)}$"), back_field),
                MessageHandler(filters.Regex(f"^{re.escape(CLEAR_BUTTON)}$"), clear),
                MessageHandler(filters.Regex(f"^{re.escape(CANCEL_BUTTON)}$"), cancel),
                MessageHandler(filters.TEXT & ~filters.COMMAND, collect_field),
            ],
            ASK_LOGO: [
                MessageHandler(filters.Regex(f"^{re.escape(BACK_BUTTON)}$"), back_from_logo),
                MessageHandler(filters.Regex(f"^{re.escape(CLEAR_BUTTON)}$"), clear),
                MessageHandler(filters.Regex(f"^{re.escape(CANCEL_BUTTON)}$"), cancel),
                MessageHandler((filters.PHOTO | filters.Document.ALL) & ~filters.COMMAND, collect_logo),
                MessageHandler(filters.ALL & ~filters.COMMAND, ask_logo_again),
            ],
        },
        fallbacks=[
            CommandHandler("start", start),
            CommandHandler("cancel", cancel),
            MessageHandler(filters.Regex(f"^{re.escape(CLEAR_BUTTON)}$"), clear),
        ],
        allow_reentry=True,
    )

    app.add_handler(conversation)
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("cancel", cancel))
    app.add_handler(MessageHandler(filters.Regex(f"^{re.escape(CLEAR_BUTTON)}$"), clear))
    return app


def main() -> None:
    try:
        asyncio.get_event_loop()
    except RuntimeError:
        asyncio.set_event_loop(asyncio.new_event_loop())

    app = build_application()
    if WEBHOOK_URL:
        app.run_webhook(
            listen="0.0.0.0",
            port=PORT,
            url_path=BOT_TOKEN,
            webhook_url=f"{WEBHOOK_URL}/{BOT_TOKEN}",
        )
    else:
        app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
